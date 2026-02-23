#!/usr/bin/env python3
"""Generate synthetic positive/negative PII test fixtures for scanner validation."""

from __future__ import annotations

import json
import shutil
import textwrap
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - generator fallback
    Document = None

try:
    from PIL import Image, ImageDraw
except Exception:  # pragma: no cover - generator fallback
    Image = None
    ImageDraw = None


ROOT = Path(__file__).resolve().parent
SUITE_ROOT = ROOT / "pii_scan_suite"


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    normalized = textwrap.dedent(content).strip() + "\n"
    path.write_text(normalized, encoding="utf-8")


def write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def verhoeff_validate(number: str) -> bool:
    mul = [
        [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
        [1, 2, 3, 4, 0, 6, 7, 8, 9, 5],
        [2, 3, 4, 0, 1, 7, 8, 9, 5, 6],
        [3, 4, 0, 1, 2, 8, 9, 5, 6, 7],
        [4, 0, 1, 2, 3, 9, 5, 6, 7, 8],
        [5, 9, 8, 7, 6, 0, 4, 3, 2, 1],
        [6, 5, 9, 8, 7, 1, 0, 4, 3, 2],
        [7, 6, 5, 9, 8, 2, 1, 0, 4, 3],
        [8, 7, 6, 5, 9, 3, 2, 1, 0, 4],
        [9, 8, 7, 6, 5, 4, 3, 2, 1, 0],
    ]
    perm = [
        [0, 1, 2, 3, 4, 5, 6, 7, 8, 9],
        [1, 5, 7, 6, 2, 8, 3, 0, 9, 4],
        [5, 8, 0, 3, 7, 9, 6, 1, 4, 2],
        [8, 9, 1, 6, 0, 4, 3, 5, 2, 7],
        [9, 4, 5, 3, 1, 2, 6, 8, 7, 0],
        [4, 2, 8, 6, 5, 7, 3, 9, 0, 1],
        [2, 7, 9, 3, 8, 0, 6, 4, 1, 5],
        [7, 0, 4, 6, 9, 1, 3, 2, 5, 8],
    ]

    checksum = 0
    for idx, digit in enumerate(reversed(number)):
        checksum = mul[checksum][perm[idx % 8][int(digit)]]
    return checksum == 0


def generate_valid_aadhaar(base11: str) -> str:
    if len(base11) != 11 or not base11.isdigit():
        raise ValueError("base11 must be exactly 11 digits")

    for check_digit in "0123456789":
        candidate = base11 + check_digit
        if verhoeff_validate(candidate):
            return candidate

    raise RuntimeError("Could not generate valid Aadhaar from base")


def format_aadhaar(number: str) -> str:
    return f"{number[:4]} {number[4:8]} {number[8:]}"


def escape_pdf_text(line: str) -> str:
    return line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_simple_text_pdf(path: Path, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    stream_lines = ["BT", "/F1 12 Tf", "72 760 Td"]
    for index, raw_line in enumerate(lines):
        safe_line = escape_pdf_text(raw_line)
        if index > 0:
            stream_lines.append("T*")
        stream_lines.append(f"({safe_line}) Tj")
    stream_lines.append("ET")

    stream_data = ("\n".join(stream_lines) + "\n").encode("latin-1", errors="replace")

    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
    )
    objects.append(
        b"<< /Length " + str(len(stream_data)).encode("ascii") + b" >>\n"
        b"stream\n" + stream_data + b"endstream"
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    out = bytearray(header)
    offsets = [0]

    for idx, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{idx} 0 obj\n".encode("ascii"))
        out.extend(obj)
        if not obj.endswith(b"\n"):
            out.extend(b"\n")
        out.extend(b"endobj\n")

    xref_offset = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    out.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        out.extend(f"{offset:010d} 00000 n \n".encode("ascii"))

    out.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )

    path.write_bytes(bytes(out))


def write_docx(path: Path, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if Document is None:
        write_text(path.with_suffix(".txt"), "DOCX generation skipped: python-docx unavailable")
        return

    document = Document()
    document.add_heading("Synthetic PII Fixture", level=1)
    for para in paragraphs:
        document.add_paragraph(para)
    document.save(path)


def write_png_text(path: Path, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if Image is None or ImageDraw is None:
        return

    image = Image.new("RGB", (1200, 500), "white")
    drawer = ImageDraw.Draw(image)
    y_pos = 20
    for line in lines:
        drawer.text((20, y_pos), line, fill="black")
        y_pos += 32
    image.save(path)


def main() -> None:
    if SUITE_ROOT.exists():
        shutil.rmtree(SUITE_ROOT)
    SUITE_ROOT.mkdir(parents=True, exist_ok=True)

    valid_aadhaar_1 = generate_valid_aadhaar("23456789012")
    valid_aadhaar_2 = generate_valid_aadhaar("34567890123")
    invalid_aadhaar = "234567890123"

    positive_txt = f"""
    Customer Intake Record
    ----------------------
    Full Name: Aarav Sharma
    Contact Email: aarav.sharma@example.in
    Contact Phone: +91 9876543210
    Aadhaar Number: {format_aadhaar(valid_aadhaar_1)}
    PAN: ABCDE1234F
    Passport: K1234567
    UPI VPA: aarav.sharma@upi
    IFSC: HDFC0001234
    Beneficiary Account Number: 123456789012
    Home Address: 22 MG Road, Bengaluru, Karnataka
    """

    write_text(SUITE_ROOT / "positive/txt/customer_profile.txt", positive_txt)

    write_text(
        SUITE_ROOT / "positive/csv/payroll_records.csv",
        """employee_id,name,email,phone,pan,aadhaar,ifsc,upi,account
E001,Aarav Sharma,aarav.sharma@example.in,+919876543210,ABCDE1234F,{a1},HDFC0001234,aarav.sharma@upi,123456789012
E002,Meera Nair,meera.nair@example.in,+919812345678,PQRST2345L,{a2},SBIN0004321,meera.nair@ybl,222233334444
""".format(a1=valid_aadhaar_1, a2=valid_aadhaar_2),
    )

    write_json(
        SUITE_ROOT / "positive/json/user_export.json",
        {
            "tenant": "demo-bank",
            "users": [
                {
                    "name": "Aarav Sharma",
                    "email": "aarav.sharma@example.in",
                    "mobile": "+91 9876543210",
                    "pan": "ABCDE1234F",
                    "aadhaar": valid_aadhaar_1,
                    "upi": "aarav.sharma@upi",
                    "ifsc": "HDFC0001234",
                    "account": "123456789012",
                }
            ],
        },
    )

    write_text(
        SUITE_ROOT / "positive/log/app_access.log",
        f"""
        2026-02-22T10:01:11Z INFO user=Aarav Sharma email=aarav.sharma@example.in action=login ip=10.0.1.2
        2026-02-22T10:05:42Z INFO payment upi=aarav.sharma@upi ifsc=HDFC0001234 account=123456789012
        2026-02-22T10:11:08Z INFO identity pan=ABCDE1234F aadhaar={valid_aadhaar_1}
        """,
    )

    write_text(
        SUITE_ROOT / "positive/md/incident_notes.md",
        f"""
        # Customer Support Note

        Reporter: Aarav Sharma
        Contact: aarav.sharma@example.in
        Backup contact: +91 9876543210
        UPI in dispute: aarav.sharma@upi
        Government IDs: PAN ABCDE1234F, Passport K1234567, Aadhaar {format_aadhaar(valid_aadhaar_1)}
        """,
    )

    write_text(
        SUITE_ROOT / "positive/xml/beneficiary.xml",
        f"""
        <beneficiary>
          <name>Aarav Sharma</name>
          <email>aarav.sharma@example.in</email>
          <phone>+91 9876543210</phone>
          <pan>ABCDE1234F</pan>
          <aadhaar>{valid_aadhaar_1}</aadhaar>
          <ifsc>HDFC0001234</ifsc>
          <account>123456789012</account>
        </beneficiary>
        """,
    )

    write_text(
        SUITE_ROOT / "positive/yaml/payment.yml",
        f"""
        customer:
          name: Aarav Sharma
          email: aarav.sharma@example.in
          phone: "+91 9876543210"
          pan: ABCDE1234F
          aadhaar: "{valid_aadhaar_1}"
        payout:
          upi: aarav.sharma@upi
          ifsc: HDFC0001234
          account: "123456789012"
        """,
    )

    write_text(
        SUITE_ROOT / "positive/yml/onboarding.yml",
        f"""
        applicant: Meera Nair
        contact_email: meera.nair@example.in
        contact_phone: +91 9812345678
        pan: PQRST2345L
        aadhaar: {valid_aadhaar_2}
        passport: M2345678
        """,
    )

    write_docx(
        SUITE_ROOT / "positive/docx/onboarding_form.docx",
        [
            "Candidate Name: Aarav Sharma",
            "Email: aarav.sharma@example.in",
            "Phone: +91 9876543210",
            f"Aadhaar: {format_aadhaar(valid_aadhaar_1)}",
            "PAN: ABCDE1234F",
            "UPI: aarav.sharma@upi",
            "IFSC: HDFC0001234",
            "Account: 123456789012",
        ],
    )

    write_simple_text_pdf(
        SUITE_ROOT / "positive/pdf/loan_application.pdf",
        [
            "Loan Applicant: Aarav Sharma",
            "Email: aarav.sharma@example.in",
            "Phone: +91 9876543210",
            f"Aadhaar: {format_aadhaar(valid_aadhaar_1)}",
            "PAN: ABCDE1234F",
            "UPI: aarav.sharma@upi",
            "IFSC: HDFC0001234",
            "Account: 123456789012",
        ],
    )

    write_png_text(
        SUITE_ROOT / "positive/images/id_card_like.png",
        [
            "Synthetic ID Card",
            "Name: Aarav Sharma",
            f"Aadhaar: {format_aadhaar(valid_aadhaar_1)}",
            "PAN: ABCDE1234F",
            "UPI: aarav.sharma@upi",
        ],
    )

    write_text(
        SUITE_ROOT / "negative/txt/no_pii_notes.txt",
        """
        Engineering weekly notes:
        - Service uptime stayed above target.
        - Latency improved after cache tuning.
        - Next sprint focuses on batch processing reliability.
        """,
    )

    write_text(
        SUITE_ROOT / "negative/csv/metrics.csv",
        """service,region,cpu_percent,memory_mb,error_rate
api,ap-south,57,640,0.12
worker,ap-south,43,720,0.08
scheduler,eu-west,38,510,0.03
""",
    )

    write_json(
        SUITE_ROOT / "negative/json/system_config.json",
        {
            "service": "pii-scanner-demo",
            "log_level": "INFO",
            "retries": 3,
            "features": {
                "tokenization_enabled": True,
                "bulk_mode": False,
            },
        },
    )

    write_text(
        SUITE_ROOT / "negative/log/healthcheck.log",
        """
        2026-02-22T09:00:00Z INFO startup complete
        2026-02-22T09:01:00Z INFO heartbeat ok
        2026-02-22T09:02:00Z INFO queue depth=7
        """,
    )

    write_text(
        SUITE_ROOT / "negative/md/public_readme.md",
        """
        # Public Product Overview

        This document contains only public feature descriptions and no customer records.
        """,
    )

    write_text(
        SUITE_ROOT / "negative/xml/catalog.xml",
        """
        <catalog>
          <item id="A1">Compute Plan</item>
          <item id="B2">Storage Plan</item>
          <item id="C3">Support Plan</item>
        </catalog>
        """,
    )

    write_text(
        SUITE_ROOT / "negative/yaml/non_sensitive.yml",
        """
        build:
          version: 1.2.3
          branch: main
          optimized: true
        deployment:
          region: ap-south
          replicas: 4
        """,
    )

    write_docx(
        SUITE_ROOT / "negative/docx/project_summary.docx",
        [
            "Project Summary",
            "This file intentionally avoids personal or sensitive identifiers.",
            "Milestone status: on track.",
        ],
    )

    write_simple_text_pdf(
        SUITE_ROOT / "negative/pdf/training_material.pdf",
        [
            "Operational Playbook",
            "This training material contains no personal identifiers.",
            "Topics: alerting, observability, rollout, rollback.",
        ],
    )

    write_png_text(
        SUITE_ROOT / "negative/images/landscape.png",
        [
            "System Diagram",
            "Gateway -> Queue -> Worker",
            "No personal data present",
        ],
    )

    write_text(
        SUITE_ROOT / "edge_cases/near_matches.txt",
        f"""
        Near matches that should not trigger custom sensitive entities:
        - Invalid Aadhaar checksum: {invalid_aadhaar}
        - PAN-like but invalid: ABCD1234F, ABCDE12345
        - UPI-like but not valid domain list: name(at)bank
        - IFSC-like but invalid: HDFC001234
        - Phone-like but invalid: +91 5123456789
        """,
    )

    write_text(
        SUITE_ROOT / "edge_cases/repeated_values.txt",
        f"""
        Repeated values for volume testing:
        aarav.sharma@example.in
        aarav.sharma@example.in
        aarav.sharma@example.in
        PAN ABCDE1234F appears twice: ABCDE1234F
        Aadhaar {format_aadhaar(valid_aadhaar_1)} repeated {format_aadhaar(valid_aadhaar_1)}
        """,
    )

    large_text_line = "This is a long non-sensitive line for file size and scanning throughput tests.\n"
    write_text(SUITE_ROOT / "edge_cases/large_file/large_non_pii.txt", large_text_line * 18000)

    binary_path = SUITE_ROOT / "edge_cases/binary/random_payload.bin"
    binary_path.parent.mkdir(parents=True, exist_ok=True)
    binary_path.write_bytes(bytes(range(256)) * 8)

    ignored_path = SUITE_ROOT / "edge_cases/node_modules/ignored_dependency.txt"
    write_text(
        ignored_path,
        f"This file has PII-like values and should be skipped when exclude_dirs includes node_modules. PAN ABCDE1234F Aadhaar {valid_aadhaar_1}",
    )

    write_json(
        SUITE_ROOT / "scan_config_text_and_docs.json",
        {
            "scan": {
                "input_paths": [str(SUITE_ROOT)],
                "recursive": True,
                "include_extensions": [
                    ".txt",
                    ".csv",
                    ".json",
                    ".log",
                    ".md",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".docx",
                    ".pdf",
                ],
                "exclude_dirs": [".git", ".idea", "venv", ".venv", "node_modules", "dist", "build", "__pycache__"],
                "exclude_file_globs": ["*.pyc", "*.pyo", "*.DS_Store"],
                "max_file_size_mb": 25,
                "ocr_images": False,
            },
            "presidio": {
                "language": "en",
                "supported_languages": ["en"],
                "nlp_engine_name": "spacy",
                "model_name": "en_core_web_lg",
                "entities": [
                    "IN_AADHAAR",
                    "IN_PAN",
                    "IN_IFSC",
                    "IN_UPI_ID",
                    "IN_PASSPORT",
                    "IN_BANK_ACCOUNT",
                    "EMAIL_ADDRESS",
                    "PHONE_NUMBER",
                    "PERSON",
                    "LOCATION",
                    "IN_ADDRESS",
                    "CREDIT_CARD",
                    "IBAN_CODE",
                    "IP_ADDRESS",
                ],
                "score_threshold": 0.35,
                "entity_score_thresholds": {
                    "PERSON": 0.6,
                    "LOCATION": 0.55,
                    "PHONE_NUMBER": 0.55,
                    "EMAIL_ADDRESS": 0.6,
                    "IN_BANK_ACCOUNT": 0.45,
                },
                "return_decision_process": False,
                "allow_list": [],
                "allow_list_match": "exact",
                "context_words": ["customer", "identity", "bank", "payment", "contact"],
                "spacy_max_length": 3000000,
                "chunk_size_chars": 200000,
                "chunk_overlap_chars": 500,
                "context_enhancer": {
                    "enabled": True,
                    "context_similarity_factor": 0.35,
                    "min_score_with_context_similarity": 0.45,
                    "context_prefix_count": 8,
                    "context_suffix_count": 2,
                },
            },
            "custom_recognizers": {
                "enable_indian_identifiers": True,
                "aadhaar_checksum_validation": True,
                "upi_generic_pattern": False,
                "upi_handle_domains": [
                    "upi",
                    "ybl",
                    "ibl",
                    "axl",
                    "paytm",
                    "okhdfcbank",
                    "okicici",
                    "oksbi",
                    "okaxis",
                ],
            },
            "rule_engine": {
                "enabled": True,
                "region": "india",
                "environment_variable": "DPDP_RULES_ENV",
                "default_environment": "default",
                "environment": "default",
                "base_rules_file": "../../config/pii_rules/india/base_rules.json",
                "environment_rules": {
                    "default": "../../config/pii_rules/india/default_rules.json",
                    "dev": "../../config/pii_rules/india/dev_rules.json",
                    "qa": "../../config/pii_rules/india/qa_rules.json",
                    "prod": "../../config/pii_rules/india/prod_rules.json",
                },
            },
            "output": {
                "path": str(SUITE_ROOT / "scan_output_example.json"),
                "pretty": True,
                "include_text_snippet": True,
                "snippet_context_chars": 24,
                "include_analysis_explanation": False,
                "include_file_hash": True,
                "mask_file_paths": False,
                "file_path_mask_mode": "full",
                "file_path_base_dir": "",
                "file_path_hash_salt": "",
            },
        },
    )

    write_json(
        SUITE_ROOT / "manifest.json",
        {
            "dataset": "pii_scan_suite",
            "description": "Synthetic positive/negative fixtures for validating scanner behavior.",
            "generated_at": utc_now(),
            "synthetic_ids": {
                "valid_aadhaar_1": valid_aadhaar_1,
                "valid_aadhaar_2": valid_aadhaar_2,
                "invalid_aadhaar": invalid_aadhaar,
                "pan_samples": ["ABCDE1234F", "PQRST2345L"],
                "ifsc_samples": ["HDFC0001234", "SBIN0004321"],
                "upi_samples": ["aarav.sharma@upi", "meera.nair@ybl"],
            },
            "expected_behavior": {
                "positive": "Should detect multiple custom entities (IN_AADHAAR, IN_PAN, IN_IFSC, IN_UPI_ID, IN_PASSPORT, IN_BANK_ACCOUNT) plus Presidio built-ins.",
                "negative": "Should have zero or very low sensitive findings; may still include generic Presidio entities depending on NLP model behavior.",
                "edge_cases": "Should validate checksum and format constraints, handle duplicates and skipped directories/files as configured.",
            },
            "folders": {
                "positive": [
                    "txt", "csv", "json", "log", "md", "xml", "yaml", "yml", "docx", "pdf", "images"
                ],
                "negative": [
                    "txt", "csv", "json", "log", "md", "xml", "yaml", "docx", "pdf", "images"
                ],
                "edge_cases": [
                    "near_matches", "repeated_values", "large_file", "binary", "node_modules"
                ],
            },
            "notes": [
                "All data in this suite is synthetic and created for scanner testing only.",
                "Image files are included for OCR scenarios; enable ocr_images and install Tesseract to scan them.",
                "node_modules fixture is intended to be skipped by default exclude_dirs.",
            ],
        },
    )

    write_text(
        SUITE_ROOT / "README.md",
        """
        # PII Scan Test Suite

        This directory contains synthetic fixtures for validating the PII scanner with positive, negative, and edge-case scenarios.

        ## Structure
        - positive/: files expected to trigger multiple entities
        - negative/: files intended to contain no sensitive identifiers
        - edge_cases/: near-matches, repeated values, large files, binary payloads, and skipped directory tests

        ## Quick Run
        From repository root:

        python main.py --config test_data/pii_scan_suite/scan_config_text_and_docs.json

        ## Notes
        - All identifiers in this dataset are synthetic test values.
        - OCR image fixtures are included but not scanned unless `ocr_images` is enabled and Tesseract is installed.
        """,
    )


if __name__ == "__main__":
    main()
